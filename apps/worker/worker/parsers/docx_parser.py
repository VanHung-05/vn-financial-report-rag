from pathlib import Path

from docx import Document as DocxDocument


def parse_docx(file_path: str | Path) -> list[dict]:
    """Extract text and tables from DOCX file.

    Returns [{"page": 1, "text": str, "tables": list[list[list]]}]
    """
    file_path = Path(file_path)
    doc = DocxDocument(str(file_path))

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)

    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            tables.append(rows)

    return [{"page": 1, "text": text, "tables": tables}]
